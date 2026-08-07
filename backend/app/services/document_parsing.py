import base64
import io

from docx import Document

from app.config import OPENAI_PDF_DETAIL, MAX_DOCUMENT_CHARS

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    DOCX_CONTENT_TYPE,
}

def _table_to_markdown(rows):
    rows = [row for row in rows if row]
    if not rows:
        return ""
    def format_row(row):
        cells = [str(cell).strip() if cell is not None else "" for cell in row]
        return "| " + " | ".join(cells) + " |"
    lines = [format_row(rows[0]), "| " + " | ".join(["---"] * len(rows[0])) + " |"]
    lines.extend(format_row(row) for row in rows[1:])
    return "\n".join(lines)

# docx is a structured XML format (not a rendering-only format like PDF), so
# extracting its text/tables locally is reliable and there's no accuracy
# tradeoff to weigh here the way there was for PDF.
def extract_docx_text(contents: bytes):
    doc = Document(io.BytesIO(contents))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        markdown = _table_to_markdown(rows)
        if markdown:
            parts.append(f"[Table]\n{markdown}")
    return "\n\n".join(parts)

def truncate_document_text(text):
    if len(text) <= MAX_DOCUMENT_CHARS:
        return text

    return (
        text[:MAX_DOCUMENT_CHARS]
        + "\n\n[Document truncated because it exceeded the configured prompt size safety limit.]"
    )

def make_pdf_input_content(filename, contents):
    base64_pdf = base64.b64encode(contents).decode("utf-8")
    return [
        {
            "type": "input_file",
            "filename": filename or "syllabus.pdf",
            "file_data": f"data:application/pdf;base64,{base64_pdf}",
            "detail": OPENAI_PDF_DETAIL,
        },
        {
            "type": "input_text",
            "text": (
                "Read this syllabus PDF and extract structured calendar/task data. "
                "Use both the PDF text and page images. If the PDF includes scans, tables, "
                "or unusual formatting, inspect the visible page content rather than returning an error."
            ),
        },
    ]

def make_text_input_content(document_text):
    return truncate_document_text(document_text)
