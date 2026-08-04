# source .venv/bin/activate
# python -m pip install -r requirements.txt
# uvicorn main:app --reload

from fastapi import FastAPI, UploadFile, File, Body
from fastapi.middleware.cors import CORSMiddleware

import os
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(Path(__file__).with_name(".env"))

from openai import OpenAI

import base64
import io
import re
from typing import Literal

from docx import Document
from pydantic import BaseModel, Field

from fastapi.responses import RedirectResponse
from fastapi.responses import HTMLResponse
from fastapi.responses import JSONResponse
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from fastapi import Request

os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1" # this is for local testing, delete before putting on the cloud

openai_api_key = os.getenv("OPENAI_API_KEY")
# CHANGE: Keep production extraction on one default model. These env vars are
# the only knobs needed to swap models or PDF rendering detail later.
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.4-nano")
OPENAI_REASONING_EFFORT = os.getenv("OPENAI_REASONING_EFFORT")
OPENAI_PDF_DETAIL = os.getenv("OPENAI_PDF_DETAIL", "auto")
MAX_DOCUMENT_CHARS = int(os.getenv("MAX_DOCUMENT_CHARS", "600000"))

SCOPES = [
    "https://www.googleapis.com/auth/calendar.events",
    "https://www.googleapis.com/auth/tasks"
]

app = FastAPI()

# Allows your React frontend to talk to this backend
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r"http://(localhost|127\.0\.0\.1):517[0-9]",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def home():
    return {"message": "Hello from the backend!"}

@app.get("/test")
def test():
    return {
        "status": "success",
        "message": "React successfully connected to FastAPI"
    }

@app.get("/auth/google")
def google_auth():
    flow = Flow.from_client_secrets_file(
        "credentials.json",
        scopes=SCOPES,
        redirect_uri="http://localhost:8000/auth/google/callback",
        autogenerate_code_verifier=False
    )

    authorization_url, state = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true",
        prompt="consent"
    )

    return RedirectResponse(authorization_url)

@app.get("/auth/google/callback")
def google_auth_callback(request: Request):
    flow = Flow.from_client_secrets_file(
        "credentials.json",
        scopes=SCOPES,
        redirect_uri="http://localhost:8000/auth/google/callback",
        autogenerate_code_verifier=False
    )

    flow.fetch_token(authorization_response=str(request.url))

    credentials = flow.credentials

    with open("token.json", "w") as token_file:
        token_file.write(credentials.to_json())

        return HTMLResponse(
                content="""
                <!doctype html>
                <html>
                    <body>
                        <script>
                            if (window.opener) {
                                window.opener.postMessage({ type: 'google-auth-success' }, 'http://localhost:5173');
                                window.close();
                            } else {
                                document.body.textContent = 'Google account connected successfully. You can close this window.';
                            }
                        </script>
                    </body>
                </html>
                """
        )

@app.get("/test-calendar")
def test_calendar():
    creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    service = build("calendar", "v3", credentials=creds)

    event = {
        "summary": "Syllabus Calendar Test Event",
        "description": "This is a test event created by my syllabus calendar app.",
        "start": {
            "dateTime": "2026-06-18T10:00:00",
            "timeZone": "America/New_York",
        },
        "end": {
            "dateTime": "2026-06-18T10:30:00",
            "timeZone": "America/New_York",
        },
    }
    
    service.events().insert(calendarId="primary", body=event).execute()
    return {"message": "Event added"}

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    DOCX_CONTENT_TYPE,
}

Confidence = Literal["high", "medium", "low"]

# Field descriptions are sent to the model as part of the Structured Outputs
# schema, so per-field mechanics live here rather than in the prompt. The
# prompt is left to explain only which bucket an item belongs in.
DATE = "ISO YYYY-MM-DD, or null if the syllabus does not state one."
TIME = "24-hour HH:MM, or null if the syllabus does not state one."
SOURCE_TEXT = "A short phrase copied verbatim from the syllabus supporting this item."
CONFIDENCE_NOTE = "Use medium or low when any part of this item was inferred."
# description is written into the Google Calendar event body / Google Task notes, so
# it is read by the student long after the upload. It is the one field that must carry
# no trace of the extraction process — source_text already holds the quote, confidence
# and missing_information already carry uncertainty.
DESCRIPTION = (
    "One short sentence saying plainly what this is, shown to the user in their "
    "calendar. Never quote the syllabus, never mention what was or was not specified, "
    "and never describe what you could or could not determine while reading. Leave it "
    "empty if the syllabus adds nothing beyond the title."
)

# CHANGE: Structured Outputs schema replaces hand-parsed JSON text so the
# model response must match the shape the review UI expects.
class Course(BaseModel):
    course_name: str | None
    course_code: str | None
    instructor: str | None
    term: str | None

class ClassMeeting(BaseModel):
    title: str = Field(
        description="The meeting type only, such as 'Lecture', 'Lab', or 'Recitation'. "
        "Do not include the course code; it is added automatically."
    )
    days_of_week: list[Literal[
        "Monday",
        "Tuesday",
        "Wednesday",
        "Thursday",
        "Friday",
        "Saturday",
        "Sunday",
    ]] = Field(
        description="Every day this meeting recurs, as full day names. A class meeting "
        "MW 10:00-11:15 is one meeting with both Monday and Wednesday here, not a "
        "separate meeting per day."
    )
    start_time: str | None = Field(description=TIME)
    end_time: str | None = Field(description=TIME)
    location: str | None
    start_date: str | None = Field(description=f"{DATE} The first date this meeting occurs.")
    end_date: str | None = Field(description=f"{DATE} The last date this meeting occurs.")
    confidence: Confidence = Field(description=CONFIDENCE_NOTE)
    source_text: str = Field(description=SOURCE_TEXT)

class ClassSchedule(BaseModel):
    found: bool
    notes: str
    meetings: list[ClassMeeting]

class ClassCancellation(BaseModel):
    title: str = Field(
        description="A short label for the cancelled session, normally 'No Class'. "
        "Not a sentence — the explanation belongs in reason and description."
    )
    date: str | None = Field(description=DATE)
    reason: str | None = Field(
        description="The cause, as a short noun phrase such as 'Spring Break'."
    )
    description: str = Field(description=DESCRIPTION)
    confidence: Confidence = Field(description=CONFIDENCE_NOTE)
    source_text: str = Field(description=SOURCE_TEXT)

# One-off things that occupy a block of time on a calendar. event_type carries
# the detail that used to be split across separate top-level arrays.
class Event(BaseModel):
    title: str = Field(
        description="The event name only, such as 'Midterm 1'. Do not include the "
        "course code; it is added automatically."
    )
    event_type: Literal[
        "exam",
        "quiz",
        "final_exam",
        "presentation",
        "review_session",
        "special_class",
        "other",
    ] = Field(
        description="What kind of event this is. Use 'other' only when the event belongs "
        "to this course and none of the named types fit."
    )
    date: str | None = Field(description=DATE)
    start_time: str | None = Field(
        description=f"{TIME} Leave null if the event falls on a regular class day "
        "with no separately stated time."
    )
    end_time: str | None = Field(description=TIME)
    location: str | None
    description: str = Field(description=DESCRIPTION)
    confidence: Confidence = Field(description=CONFIDENCE_NOTE)
    source_text: str = Field(description=SOURCE_TEXT)

# Things with a deadline rather than a duration. task_type carries the detail.
class Task(BaseModel):
    title: str = Field(
        description="The task name only, such as 'Project 1'. Do not include the "
        "course code; it is added automatically."
    )
    task_type: Literal[
        "homework",
        "assignment",
        "paper",
        "project",
        "lab",
        "problem_set",
        "other",
    ] = Field(
        description="What kind of work this is. Use 'project' only when the syllabus "
        "itself calls it a project. Use 'other' if none fit."
    )
    due_date: str | None = Field(
        description=f"{DATE} If the item is listed against a week range rather than a "
        "single day, use the last day of that range and set confidence to medium."
    )
    due_time: str | None = Field(description=TIME)
    description: str = Field(description=DESCRIPTION)
    confidence: Confidence = Field(description=CONFIDENCE_NOTE)
    source_text: str = Field(description=SOURCE_TEXT)

class Reading(BaseModel):
    title: str = Field(
        description="The reading name only. Do not include the course code; it is "
        "added automatically."
    )
    reading_type: Literal["textbook", "article", "class_notes", "book", "other"]
    due_date: str | None = Field(
        description=f"{DATE} If the reading is listed against a week range rather than "
        "a single day, use the last day of that range and set confidence to medium."
    )
    due_time: str | None = Field(description=TIME)
    description: str = Field(
        description=DESCRIPTION + " For a reading, name the topic it covers rather than "
        "the date it is assigned, which is already in due_date."
    )
    confidence: Confidence = Field(description=CONFIDENCE_NOTE)
    source_text: str = Field(description=SOURCE_TEXT)

class SyllabusExtraction(BaseModel):
    course: Course
    class_schedule: ClassSchedule
    class_cancellations: list[ClassCancellation]
    events: list[Event]
    tasks: list[Task]
    readings: list[Reading]
    missing_information: list[str]
    warnings: list[str]

# --- Deduplication -----------------------------------------------------------------
#
# Syllabi name the same thing in more than one place. Psyc lists every essay twice:
# once in the weekly schedule ("Essay 2") and again in a separate essay schedule
# ("Unplug Day Essay"). The model dutifully records both, so the user sees six essays
# where there are three. Prompt rules catch this only sometimes; these three do it every
# time. Two items are the same thing when:
#
#   1. their titles match, or one is the start of the other
#      ("Principles of Life (Hillis et al.)" and "...(Hillis et al.) a")
#   2. they fall on the same date, are the same type, and one is a generic label
#      ("Essay 2" next to "Unplug Day Essay", both due 3/30)
#   3. the same title appears in both events and tasks - an exam belongs in events
#
# Rule 2 requires a generic label on one side deliberately: two differently-named
# assignments due the same day are two assignments, and must not be merged.

# A title that names a category and a number but nothing specific about the work.
GENERIC_TITLE = re.compile(
    r"^(essay|assignment|homework|hw|quiz|test|exam|midterm|final|project|paper"
    r"|problem set|pset|lab|reading|questionnaire)\s*\d*$"
)

def normalized_title(title):
    """Lowercased, punctuation stripped, so 'Essay 1.' and 'essay 1' compare equal."""
    return " ".join(re.sub(r"[^a-z0-9]+", " ", (title or "").lower()).split())

def same_item(a, b, date_field, type_field):
    title_a = normalized_title(a.get("title"))
    title_b = normalized_title(b.get("title"))
    date = a.get(date_field)

    # Same day, or both undated. Two "No Class" cancellations a week apart are two
    # separate days off, however identical their titles.
    if date != b.get(date_field):
        return False

    # The shorter title must end on a word boundary in the longer one, or "Project 1"
    # would swallow "Project 10".
    if title_a and title_b:
        short, long = sorted((title_a, title_b), key=len)
        if long == short or long.startswith(short + " "):
            return True

    if not date:
        return False
    if type_field and a.get(type_field) != b.get(type_field):
        return False

    # The generic label must name the same kind of work as the title it is merging
    # into - "Essay 1" into "Personal Application Essay". Without this, "Questionnaire 1"
    # merges with "Information Sheet and Written Release Form" purely for sharing a date,
    # which loses a real item.
    for generic, other in ((title_a, title_b), (title_b, title_a)):
        match = GENERIC_TITLE.match(generic)
        if match and match.group(1) in other:
            return True
    return False

def dedupe_list(items, date_field, type_field):
    """Keep one of each item, preferring the longer - and so more descriptive - title."""
    kept = []
    for item in items:
        match = next((k for k in kept if same_item(item, k, date_field, type_field)), None)
        if match is None:
            kept.append(item)
        elif len(item.get("title") or "") > len(match.get("title") or ""):
            kept[kept.index(match)] = item
    return kept

CONFIDENCE_RANK = {"high": 0, "medium": 1, "low": 2}
WEEKDAY_ORDER = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday",
                 "Saturday", "Sunday"]

def compatible(a, b):
    """Two values describe the same thing when they agree or one is simply absent."""
    return not a or not b or a == b

def same_meeting(a, b):
    """Whether two entries describe one class meeting. Matching on agreement rather
    than equality is what merges a schedule the model split in two - the syllabus
    states the lecture's days and times in one place and its date range in another, so
    neither copy is a duplicate of the other but together they are one meeting.

    One day set inside the other counts as agreement, because the other way a schedule
    splits is per day: a Monday/Wednesday lecture returned again as a bare Monday entry
    and a bare Wednesday entry. A genuinely separate session on one of those days -
    a Wednesday lab, an extra Friday hour - differs in title, room or time, and every
    one of those still has to agree below."""
    days_a = set(a.get("days_of_week") or [])
    days_b = set(b.get("days_of_week") or [])
    if not (days_a <= days_b or days_b <= days_a):
        return False
    if not compatible(a.get("start_time"), b.get("start_time")):
        return False
    if not compatible(a.get("end_time"), b.get("end_time")):
        return False

    # Location must agree: Bio runs two recitation sections at Thursday 16:50 in
    # different rooms, and they are different sections, not a duplicate.
    location_a = normalized_title(a.get("location"))
    location_b = normalized_title(b.get("location"))
    if not compatible(location_a, location_b):
        return False
    # Nothing else can be in the same room at the same time, so two known and equal
    # locations settle it - which catches the same class recorded once as "Lecture"
    # and again as "Class". Otherwise the titles have to agree too, since two
    # untimed, unplaced meetings may genuinely differ.
    if location_a and location_b:
        return True
    return compatible(normalized_title(a.get("title")), normalized_title(b.get("title")))

def merge_meetings(a, b):
    """One meeting carrying whatever either copy knew. Takes the widest date range:
    a truncated end date is the likelier mistake, and the same holds of a start date
    that begins mid-term."""
    merged = dict(a)
    for field in ("start_time", "end_time", "location", "source_text"):
        merged[field] = a.get(field) or b.get(field)
    # Keep every day either copy knew about, in the schedule's own order.
    days = list(a.get("days_of_week") or [])
    days += [d for d in (b.get("days_of_week") or []) if d not in days]
    merged["days_of_week"] = sorted(days, key=WEEKDAY_ORDER.index)
    # Prefer the more descriptive of two disagreeing titles, as dedupe_list does.
    merged["title"] = max((a.get("title") or "", b.get("title") or ""), key=len)
    merged["start_date"] = min(filter(None, (a.get("start_date"), b.get("start_date"))),
                               default=None)
    merged["end_date"] = max(filter(None, (a.get("end_date"), b.get("end_date"))),
                             default=None)
    # A merged meeting is partly assembled, so it is only as trustworthy as its
    # weaker half.
    merged["confidence"] = max(
        (a.get("confidence"), b.get("confidence")),
        key=lambda c: CONFIDENCE_RANK.get(c, 0),
    )
    return merged

def dedupe_meetings(meetings):
    """A class recurring on the same days at the same time is one meeting, however
    many times the syllabus mentions it. Psyc returns its Monday/Wednesday lecture
    twice, sometimes with different end dates."""
    kept = []
    for meeting in meetings:
        match = next((i for i, k in enumerate(kept) if same_meeting(meeting, k)), None)
        if match is None:
            kept.append(meeting)
        else:
            kept[match] = merge_meetings(kept[match], meeting)
    return kept

DEDUPE_FIELDS = {
    "events": ("date", "event_type"),
    "tasks": ("due_date", "task_type"),
    "readings": ("due_date", "reading_type"),
    "class_cancellations": ("date", None),
}

def deduplicate(payload):
    for key, (date_field, type_field) in DEDUPE_FIELDS.items():
        payload[key] = dedupe_list(payload.get(key, []), date_field, type_field)

    schedule = payload.setdefault("class_schedule", {})
    schedule["meetings"] = dedupe_meetings(schedule.get("meetings", []))

    # Rule 3. Anything recorded as both an event and a task stays an event.
    event_titles = {normalized_title(e.get("title")) for e in payload["events"]}
    payload["tasks"] = [
        t for t in payload["tasks"] if normalized_title(t.get("title")) not in event_titles
    ]
    return payload

# Every item reaching Calendar or Tasks is prefixed with the course code so it is
# identifiable out of context. Done here rather than in the prompt so it is applied
# uniformly and cannot drift with the model's phrasing.
def titled(payload, name, fallback):
    name = (name or "").strip() or fallback
    course_code = (payload.get("course") or {}).get("course_code")
    if not course_code:
        return name
    course_code = course_code.strip()
    if not course_code or name.startswith(course_code):
        return name
    return f"{course_code}: {name}"

# Applied to the extraction before it reaches the review screen, so the titles the
# user edits are the titles that get created. titled() still runs at sync time and is
# idempotent, which covers anything the user adds by hand while reviewing.
def apply_course_code(payload):
    defaults = {
        "events": "Event",
        "tasks": "Task",
        "readings": "Reading",
        "class_cancellations": "No Class",
    }
    for key, fallback in defaults.items():
        for item in payload.get(key, []):
            item["title"] = titled(payload, item.get("title"), fallback)

    meetings = payload.get("class_schedule", {}).get("meetings", [])
    for meeting in meetings:
        meeting["title"] = titled(payload, meeting.get("title"), "Class Meeting")

    return payload

def error_response(status_code, error, message):
    return JSONResponse(
        status_code=status_code,
        content={"error": error, "message": message},
    )

def _table_to_markdown(rows):
    rows = [row for row in rows if row]
    if not rows:
        return ""
    def format_row(row):
        cells = [str(cell).strip() if cell is not None else "" for cell in row]
        return "| " + " | ".join(cells) + " |"
    lines = [format_row(rows[0]), "| " + " | ".join(["---"] * len(rows[0])) + " |"]
    lines.extend(format_row(row) for row in rows[1:])
    return "\n".join(lines)

# docx is a structured XML format (not a rendering-only format like PDF), so
# extracting its text/tables locally is reliable and there's no accuracy
# tradeoff to weigh here the way there was for PDF.
def extract_docx_text(contents: bytes):
    doc = Document(io.BytesIO(contents))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        markdown = _table_to_markdown(rows)
        if markdown:
            parts.append(f"[Table]\n{markdown}")
    return "\n\n".join(parts)

def truncate_document_text(text):
    if len(text) <= MAX_DOCUMENT_CHARS:
        return text

    return (
        text[:MAX_DOCUMENT_CHARS]
        + "\n\n[Document truncated because it exceeded the configured prompt size safety limit.]"
    )

def make_pdf_input_content(filename, contents):
    base64_pdf = base64.b64encode(contents).decode("utf-8")
    return [
        {
            "type": "input_file",
            "filename": filename or "syllabus.pdf",
            "file_data": f"data:application/pdf;base64,{base64_pdf}",
            "detail": OPENAI_PDF_DETAIL,
        },
        {
            "type": "input_text",
            "text": (
                "Read this syllabus PDF and extract structured calendar/task data. "
                "Use both the PDF text and page images. If the PDF includes scans, tables, "
                "or unusual formatting, inspect the visible page content rather than returning an error."
            ),
        },
    ]

def make_text_input_content(document_text):
    return truncate_document_text(document_text)

@app.post("/upload-syllabus")
async def upload_syllabus(file: UploadFile = File(...)):
    if file.content_type not in ALLOWED_CONTENT_TYPES:
        return error_response(
            415,
            "unsupported_file_type",
            "Unsupported file type. Please upload a PDF or DOCX file.",
        )

    if not openai_api_key:
        return error_response(
            503,
            "missing_openai_api_key",
            "OPENAI_API_KEY is not configured on the backend.",
        )

    contents = await file.read()

    if file.content_type == "application/pdf":
        # CHANGE: Send PDFs directly to OpenAI so normal or scanned syllabi do
        # not get blocked by local text-extraction failures.
        input_content = make_pdf_input_content(file.filename, contents)
    elif file.content_type == DOCX_CONTENT_TYPE:
        try:
            extracted_text = extract_docx_text(contents)
        except Exception:
            return error_response(
                422,
                "docx_text_extraction_failed",
                "Could not read this DOCX file. Please check that it is not corrupted.",
            )

        if not extracted_text.strip():
            return error_response(
                422,
                "docx_has_no_extractable_text",
                "Could not find text in this DOCX file.",
            )

        document_text = f"Here is the text extracted from the syllabus document, including any tables converted to markdown:\n\n{extracted_text}"
        input_content = make_text_input_content(document_text)

    client = OpenAI(api_key=openai_api_key)

    instructions = """You are a syllabus extraction assistant for a college calendar app.

Read the full syllabus text and extract structured scheduling information.

Sort everything you find into exactly one of five lists. Every item belongs in a single
list: once you have recorded something, never record it again in another list. Ask "does
this take up a block of time, or is it a deadline?" first, then pick the list:

- class_schedule.meetings - a class that repeats every week, such as "MW 2:00-3:15" or
  "every Tuesday and Thursday".
- events - a one-off thing that occupies a block of time and belongs to this course:
  exams, midterms, finals, quizzes, presentations, review sessions, and special class
  meetings.
- tasks - work that is due by a deadline: homework, assignments, papers, projects,
  labs, and problem sets.
- readings - textbook chapters, articles, books, and class notes to read. Where a
  schedule row gives both a lecture topic and the material for it, the reading is the
  material, never the topic, and its date is that row's date.
- class_cancellations - "No Class", holidays, breaks, and university closures.

The confusions worth naming: a repeating class is never an event, an exam is never a
task because a student sits it rather than hands it in, a reading is never a task, and
a break is never an event.

Rules:
- Do not invent dates, times, titles, or locations.
- If information is missing or unclear, use null and explain it in missing_information or warnings.
- Record every class meeting you can identify, even when details are missing or look
  wrong. Set the unknown fields to null and say what was missing in missing_information.
  Never drop a meeting for being incomplete. Where a stated time is plainly a typo,
  record the time the syllabus meant and note the correction in warnings.
- Give each distinct recurring pattern its own entry, with all of that pattern's days
  grouped into its days_of_week. Lecture, lab and recitation are separate entries, and
  when one of them is offered as several sections at different times, record every
  section rather than collapsing them or picking one.
- A week-by-week or dated schedule table lists individual sessions of a meeting you have
  already recorded. Use those rows to find start_date, end_date, exams, tasks and
  readings, but never add a meeting for one.
- A table of sections is not that kind of table: rows giving a day, a time and a room but
  no calendar date. Every row there is a separate section and is its own meeting.
- A break spanning several days becomes one cancellation for each day the class would
  otherwise have met - not one entry for the whole range, and not one per calendar day.
- Only record scheduled work and dated events. Policy text - make-up rules, grading
  breakdowns, contact instructions, honor-code pledges, attendance requirements - is not
  an item, however much it describes something a student must do.
- Every item must be something this course scheduled. A date the university sets for all
  students - a break, a holiday, an add/drop or withdrawal deadline, a pass/fail or
  grading-option deadline, a registration window, reading days, commencement - goes in no
  list at all, except that a break stopping this class from meeting is a cancellation.
- Do not include an item that has no date, unless it is a repeating class meeting.
- Record one item per dated occurrence, and never one item standing in for several. A
  syllabus saying homework is assigned weekly without listing the weeks yields no tasks
  at all - not one called "Homework Assignments" - and the same goes for "Weekly Quizzes"
  or "Discussion Posts". Say so in warnings instead. Outside class_schedule.meetings,
  every item covers exactly one date.
- If a reading is listed as TBD, leave it out of readings and add it to warnings instead.
- Take start_date and end_date from the first and last dated class meetings where the
  syllabus has a dated schedule, otherwise from whatever states the term's span. Mark
  confidence as medium if inferred.
- Set class_schedule.found to false only when the syllabus says nothing at all about when
  the class meets. Finding a meeting but not its times is still found = true."""

    request = {
        "model": OPENAI_MODEL,
        "instructions": instructions,
        "input": [
            {
                "role": "user",
                "content": input_content,
            }
        ],
        "text_format": SyllabusExtraction,
        "prompt_cache_key": "syllabus-extraction-v1",
    }
    if OPENAI_REASONING_EFFORT:
        request["reasoning"] = {"effort": OPENAI_REASONING_EFFORT}

    try:
        response = client.responses.parse(**request)
    except Exception:
        return error_response(
            502,
            "openai_request_failed",
            "OpenAI could not process the extracted syllabus text. Please try again.",
        )

    parsed = response.output_parsed
    if not parsed:
        return error_response(
            502,
            "invalid_model_response",
            "Model returned an unexpected response shape.",
        )

    # Deduplicate before prefixing, so the course code does not mask a generic title.
    return apply_course_code(deduplicate(parsed.model_dump(mode="json")))

def create_repeating_event(
    service,
    name,
    start,
    end,
    repeat_days,
    repeat_until,
    location="", 
    colorId=None,
    description="",
    timezone="America/New_York"
    ):
    day_map = {
        "Monday": "MO",
        "Tuesday": "TU",
        "Wednesday": "WE",
        "Thursday": "TH",
        "Friday": "FR",
        "Saturday": "SA",
        "Sunday": "SU",

        "Mon": "MO",
        "Tue": "TU",
        "Tues": "TU",
        "Wed": "WE",
        "Thu": "TH",
        "Thur": "TH",
        "Thurs": "TH",
        "Fri": "FR",
        "Sat": "SA",
        "Sun": "SU",

        "M": "MO",
        "T": "TU",
        "W": "WE",
        "R": "TH",
        "F": "FR",
    }

    byday = ",".join(day_map[day] for day in repeat_days)
    until = repeat_until.replace("-", "")

    body = {
        "summary": name,
        "description": description,
        "start": {
            "dateTime": start,
            "timeZone": timezone,
        },
        "end": {
            "dateTime": end,
            "timeZone": timezone,
        },
        "location": location,
        "recurrence": [
            f"RRULE:FREQ=WEEKLY;BYDAY={byday};UNTIL={until}"
        ]
    }
    if colorId:
        body["colorId"] = str(colorId)
    print(body)

    service.events().insert(calendarId="primary", body=body).execute()

def create_event(
    service, 
    name, 
    start, 
    end, 
    description="", 
    location="", 
    colorId = None, 
    timezone="America/New_York",
    all_day=False
):
    if all_day:
        body = {
            "summary": name,
            "description": description,
            "start": {
                "date": start,
            },
            "end": {
                "date": end,
            },
            "location": location
        }
    else:
        body = {
            "summary": name,
            "description": description,
            "start": {
                "dateTime": start,
                "timeZone": timezone,
            },
            "end": {
                "dateTime": end,
                "timeZone": timezone,
            },
            "location": location
        }
    if colorId:
        body["colorId"] = str(colorId)
    
    service.events().insert(calendarId="primary", body=body).execute()

def create_task(
    service,
    name,
    due_date,
    due_time=None,
    description="",
    tasklist="@default"
):
    if due_time:
        due = f"{due_date}T{due_time}:00.000Z"
    else:
        due = f"{due_date}T00:00:00.000Z"

    body = {
        "title": name,
        "notes": description,
        "due": due,
    }

    service.tasks().insert(tasklist=tasklist, body=body).execute()

# A recurring meeting needs a start and end date to be scheduled at all, and the model
# supplies them unreliably. Rather than silently skipping the meeting, fall back to the
# span of everything else that is dated in the same syllabus - exams, due dates,
# cancellations - which brackets the term closely enough to be useful and is at worst a
# week or so long at the end. Returns (None, None) only if nothing anywhere has a date.
def term_bounds(payload):
    dates = []
    for meeting in payload.get("class_schedule", {}).get("meetings", []):
        dates += [meeting.get("start_date"), meeting.get("end_date")]
    for key in ("events", "class_cancellations"):
        dates += [item.get("date") for item in payload.get(key, [])]
    for key in ("tasks", "readings"):
        dates += [item.get("due_date") for item in payload.get(key, [])]

    dates = sorted(d for d in dates if d)
    if not dates:
        return None, None
    return dates[0], dates[-1]

def add_class_schedule(payload, service, color_id):
    class_schedule = payload.get("class_schedule", {})
    meetings = class_schedule.get("meetings", [])

    if not meetings:
        return

    fallback_start, fallback_end = term_bounds(payload)
    report = {"added": 0, "dates_inferred": 0, "skipped": []}

    for meeting in meetings:
        title = titled(payload, meeting.get("title"), "Class Meeting")
        days_of_week = meeting.get("days_of_week", [])
        start_time = meeting.get("start_time")
        end_time = meeting.get("end_time")
        start_date = meeting.get("start_date")
        end_date = meeting.get("end_date")
        location = meeting.get("location", "")

        inferred = not (start_date and end_date)
        start_date = start_date or fallback_start
        end_date = end_date or fallback_end

        if not (days_of_week and start_time and end_time and start_date and end_date):
            # Still unschedulable. Name it in the response rather than dropping it
            # silently, which previously made a whole schedule vanish with no error.
            report["skipped"].append(title)
            continue

        start = f"{start_date}T{start_time}:00"
        end = f"{start_date}T{end_time}:00"

        create_repeating_event(
            service,
            title,
            start,
            end,
            days_of_week,
            end_date,
            location,
            color_id
        )
        report["added"] += 1
        report["dates_inferred"] += 1 if inferred else 0

    return report

def add_events_to_calendar(payload, service, color_id):
    for event in payload.get("events", []):
        date = event.get("date")
        start_time = event.get("start_time")
        end_time = event.get("end_time")
        if not date:
            continue

        title = titled(payload, event.get("title"), "Event")

        if start_time and end_time:
            start = date + "T" + start_time + ":00"
            end = date + "T" + end_time + ":00"
            create_event(
                service,
                title,
                start,
                end,
                event.get("description", ""), # this is if we want descriptions included, if not then make this ""
                event.get("location", ""),
                color_id
            )
            continue

        next_day = date
        create_event(
            service,
            title,
            date,
            next_day,
            event.get("description", ""),
            event.get("location", ""),
            color_id,
            all_day=True,
        )

def add_due_items(payload, items, service, default_title):
    for item in items:
        due_date = item.get("due_date")
        if not due_date:
            continue

        create_task(
            service,
            titled(payload, item.get("title"), default_title),
            due_date,
            item.get("due_time"),
            item.get("description", ""),
        )

def add_tasks(payload, service):
    add_due_items(payload, payload.get("tasks", []), service, "Task")

def add_readings(payload, service):
    add_due_items(payload, payload.get("readings", []), service, "Reading")

@app.post("/add-events")
def add_events(payload: dict = Body(...)):
    print("adding events")

    if not os.path.exists("token.json"):
        return JSONResponse(
            status_code=401,
            content={"error": "not_authenticated", "message": "Google account not connected"},
        )

    creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    calendar_service = build("calendar", "v3", credentials=creds)
    tasks_service = build("tasks", "v1", credentials=creds)
    color_id = payload.get("color_id", 1)

    schedule_report = add_class_schedule(payload, calendar_service, color_id)
    add_events_to_calendar(payload, calendar_service, color_id)
    add_tasks(payload, tasks_service)
    add_readings(payload, tasks_service)
    print("added events successfully")

    # Report what happened to the class schedule so a meeting that could not be
    # scheduled is visible to the caller instead of disappearing.
    return {"message": "Events added successfully", "class_schedule": schedule_report}
